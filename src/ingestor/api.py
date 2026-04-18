# Fichier: src/ingestor/api.py
from __future__ import annotations

import hashlib
import hmac
import importlib.util
import ipaddress
import json
import logging
import mimetypes
import os
import socket
import sys
import tempfile
import time
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from pathlib import Path
from types import ModuleType
from typing import Any, Literal, cast
from urllib.parse import urlparse

import chromadb
import docx
import requests
from bs4 import BeautifulSoup
from chromadb.config import Settings
from fastapi import FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.exceptions import RequestValidationError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings import OllamaEmbeddings
from langchain_core.documents import Document
from langchain_google_community import GoogleDriveLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from prometheus_client import CONTENT_TYPE_LATEST
from pydantic import AliasChoices, BaseModel, ConfigDict, Field

try:
    from .mm_adapter import Chunk, parse_multimodal
except ImportError:
    # Allow running when the module is executed as a top-level script (e.g. inside Docker).
    from mm_adapter import Chunk, parse_multimodal  # type: ignore[no-redef]


def _load_metrics_module() -> ModuleType:
    module_name = "src.ingestor.metrics"
    existing = sys.modules.get(module_name)
    if isinstance(existing, ModuleType):
        return existing
    spec = importlib.util.spec_from_file_location(
        module_name,
        Path(__file__).with_name("metrics.py"),
    )
    if spec and spec.loader:
        module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = module
        spec.loader.exec_module(module)
        return module
    raise ImportError("Unable to load metrics module")


ingest_metrics: ModuleType = _load_metrics_module()

# --- Configuration ---
CHROMA_HOST = os.getenv("CHROMA_HOST", "localhost")
CHROMA_PORT = int(os.getenv("CHROMA_PORT", "8000"))
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
EMBED_MODEL = os.getenv("EMBED_MODEL", "nomic-embed-text")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "rag_education")

# Multi-collection routing by section
COLLECTION_MAP: dict[str, str] = {
    "education": os.getenv("COLLECTION_EDUCATION", "rag_education"),
    "web3": os.getenv("COLLECTION_WEB3", "rag_web3"),
    "blockchain": os.getenv("COLLECTION_WEB3", "rag_web3"),
}


def resolve_collection_name(
    section: str | None = None, collection: str | None = None,
) -> str:
    """Resolve target ChromaDB collection from section or explicit name."""
    if collection:
        return collection.strip()
    key = (section or "").strip().lower()
    return COLLECTION_MAP.get(key, COLLECTION_NAME)

CHROMA_REQUEST_TIMEOUT = float(os.getenv("CHROMA_REQUEST_TIMEOUT", "30"))
OLLAMA_REQUEST_TIMEOUT = float(os.getenv("OLLAMA_REQUEST_TIMEOUT", "30"))
MAX_REMOTE_BYTES = int(os.getenv("MAX_REMOTE_BYTES", str(10 * 1024 * 1024)))
LOCAL_SOURCE_ROOT = Path(
    os.getenv("LOCAL_SOURCE_ROOT", "/data/uploads")).resolve()
ALLOW_UNRESTRICTED_LOCAL = os.getenv(
    "ALLOW_UNRESTRICTED_LOCAL", "false").lower() == "true"
URL_SCHEMES_ALLOWED = {"http", "https"}

INGEST_CHUNK_SIZE = int(os.getenv("INGEST_CHUNK_SIZE", "1000"))
INGEST_CHUNK_OVERLAP = int(os.getenv("INGEST_CHUNK_OVERLAP", "150"))
METRICS_ENABLED = ingest_metrics.METRICS_ENABLED
MULTIMODAL_ENABLED = os.getenv("MULTIMODAL_ENABLED", "true").lower() == "true"
MM_PARSER_TIMEOUT = float(os.getenv("MM_PARSER_TIMEOUT", "1800"))
MM_MAX_CHARS_PER_CHUNK = int(os.getenv("MM_MAX_CHARS_PER_CHUNK", "1200"))
MM_CACHE_DIR = os.getenv("MM_CACHE_DIR", "/data/mm-cache")

# Keep metrics isolated per module import to avoid duplicate registration in tests.
METRIC_REGISTRY = ingest_metrics.REGISTRY
REQUEST_COUNT = ingest_metrics.REQUEST_COUNT
REQUEST_LATENCY = ingest_metrics.REQUEST_LATENCY
INGEST_RESULT = ingest_metrics.INGEST_RESULT
ingest_requests_total = ingest_metrics.ingest_requests_total

logger = logging.getLogger(__name__)


@dataclass
class PreparedBatch:
    ids: list[str]
    documents: list[str]
    metadatas: list[dict[str, str]]
    modality: str


MEDIA_SOURCE_TYPES = frozenset({"video", "image", "audio"})


def _dedupe_prepared_batch(prepared: PreparedBatch) -> tuple[PreparedBatch, int]:
    seen_ids: set[str] = set()
    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict[str, str]] = []
    skipped = 0

    for chunk_id, document, metadata in zip(
        prepared.ids, prepared.documents, prepared.metadatas, strict=False
    ):
        if chunk_id in seen_ids:
            skipped += 1
            continue
        seen_ids.add(chunk_id)
        ids.append(chunk_id)
        documents.append(document)
        metadatas.append(metadata)

    return PreparedBatch(ids=ids, documents=documents, metadatas=metadatas, modality=prepared.modality), skipped


def _validate_upload_mode(source_type: str, mode: str) -> None:
    normalized_source_type = (source_type or "").strip().lower()
    normalized_mode = (mode or "text").strip().lower() or "text"
    if normalized_source_type not in MEDIA_SOURCE_TYPES:
        return
    if normalized_mode != "multimodal":
        raise HTTPException(
            status_code=400,
            detail="Ce type de fichier nécessite mode=multimodal.",
        )
    if not MULTIMODAL_ENABLED:
        raise HTTPException(status_code=400, detail="Multimodal ingest disabled")

app = FastAPI(title="RAG Ingestor API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.exception_handler(RequestValidationError)
async def _validation_exception_handler(
    request: Request, exc: RequestValidationError,
) -> JSONResponse:
    """Return a clearer error when the JSON body cannot be parsed.

    This commonly happens when filenames with special Unicode characters
    (emojis, supplementary-plane chars) cause encoding issues in the
    request body.
    """
    logger.warning("Request validation error on %s: %s", request.url.path, exc)
    return JSONResponse(
        status_code=422,
        content={
            "status": "error",
            "detail": "Erreur de validation du corps de la requête. "
            "Vérifiez que les noms de fichiers ne contiennent pas "
            "de caractères spéciaux (emojis, guillemets, etc.).",
            "errors": [str(e) for e in exc.errors()[:5]],
        },
    )


@app.middleware("http")
async def _metrics_middleware(request, call_next):
    start = time.perf_counter()
    code = 500
    try:
        response = await call_next(request)
        code = getattr(response, "status_code", 500)
    except Exception:
        code = 500
        raise
    finally:
        if ingest_metrics.METRICS_ENABLED:
            elapsed = time.perf_counter() - start
            path = request.url.path
            method = request.method
            REQUEST_LATENCY.labels(path=path, method=method).observe(elapsed)
            REQUEST_COUNT.labels(path=path, method=method, code=str(code)).inc()
    return response


def _record_ingest_metrics(ok: bool) -> None:
    if ingest_metrics.METRICS_ENABLED:
        INGEST_RESULT.labels(status="ok" if ok else "fail").inc()


def _record_ingest_outcome(source: str, modality: str, status: str) -> None:
    if not METRICS_ENABLED:
        return
    safe_source = (source or "unknown").strip().lower() or "unknown"
    safe_modality = (modality or "unknown").strip().lower() or "unknown"
    safe_status = (status or "unknown").strip().lower() or "unknown"
    ingest_requests_total.labels(
        source=safe_source, modality=safe_modality, status=safe_status
    ).inc()

# --- Modèle de requête ---


class IngestRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True, extra="ignore")
    source_type: Literal["url", "gdrive_folder", "pdf", "docx", "markdown", "md", "video", "auto", "image", "audio"] = Field(
        alias="sourceType",
        validation_alias=AliasChoices("source_type", "sourceType"),
    )
    source: str = Field(
        alias="sourceUrl",
        validation_alias=AliasChoices("source", "sourceUrl"),
    )
    metadata_hints: dict[str, str] = Field(
        default_factory=dict,
        alias="metadata",
        validation_alias=AliasChoices("hints", "metadata"),
    )

# --- Utilitaires ---


def normalize_metadata(d: dict) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for key, value in d.items():
        if value in (None, ""):
            continue
        normalized[str(key).strip().lower().replace(" ", "_")] = str(value)
    return normalized


def get_content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _resolve_local_path(raw_path: str, *, allow_tmp: bool = False) -> Path:
    candidate = Path(raw_path)
    if not candidate.is_absolute():
        candidate = (LOCAL_SOURCE_ROOT / candidate).resolve()
    else:
        candidate = candidate.resolve()
    in_allowed_root = str(candidate).startswith(str(LOCAL_SOURCE_ROOT))
    in_tmp = allow_tmp and str(candidate).startswith("/tmp/")
    if not ALLOW_UNRESTRICTED_LOCAL and not in_allowed_root and not in_tmp:
        raise HTTPException(
            status_code=400, detail="Chemin local en dehors de la zone autorisée")
    if not candidate.exists():
        raise HTTPException(status_code=400, detail="Fichier introuvable")
    if not candidate.is_file():
        raise HTTPException(
            status_code=400, detail="Le chemin indiqué n'est pas un fichier")
    return candidate


def _get_client_ip(request: Any) -> str:
    headers = getattr(request, "headers", {}) or {}
    forwarded = headers.get("X-Forwarded-For") or headers.get("x-forwarded-for")
    if isinstance(forwarded, str) and forwarded.strip():
        primary = forwarded.split(",")[0].strip()
        if primary:
            return primary
    client = getattr(request, "client", None)
    host = getattr(client, "host", None)
    if isinstance(host, str) and host:
        return host
    return "127.0.0.1"


def _ip_allowed(ip_str: str, allowlist: str | None) -> bool:
    if not allowlist:
        return True
    try:
        ip_obj = ipaddress.ip_address(ip_str)
    except ValueError:
        return False
    for cidr in allowlist.split(","):
        network = cidr.strip()
        if not network:
            continue
        try:
            if ip_obj in ipaddress.ip_network(network, strict=False):
                return True
        except ValueError:
            continue
    return False


def _extract_bearer_token(headers: Any) -> str | None:
    """Extract token from Authorization: Bearer <token> header."""
    auth = headers.get("Authorization") or headers.get("authorization") or ""
    if auth.startswith("Bearer "):
        return auth[7:].strip() or None
    return None


def _enforce_security(request: Any, _req: Any) -> None:
    headers = getattr(request, "headers", {}) or {}
    token_env = os.getenv("INGESTOR_API_TOKEN") or os.getenv("INGEST_AUTH_TOKEN")
    if not token_env:
        logger.error("INGESTOR_API_TOKEN is not configured — rejecting request")
        raise HTTPException(status_code=503, detail="API token not configured on server")

    header_token = (
        headers.get("X-API-Token")
        or headers.get("x-api-token")
        or _extract_bearer_token(headers)
    )
    if not header_token or not hmac.compare_digest(header_token, token_env):
        raise HTTPException(status_code=401, detail="Unauthorized")

    allowlist = os.getenv("INGESTOR_IP_ALLOWLIST")
    if allowlist and not _ip_allowed(_get_client_ip(request), allowlist):
        raise HTTPException(status_code=403, detail="Forbidden")


def get_chroma_client() -> Any:
    timeout_seconds = max(1, int(CHROMA_REQUEST_TIMEOUT))
    settings = Settings(
        chroma_server_host=CHROMA_HOST,
        chroma_server_http_port=CHROMA_PORT,
        anonymized_telemetry=False,
        chroma_logservice_request_timeout_seconds=timeout_seconds,
        chroma_sysdb_request_timeout_seconds=timeout_seconds,
        chroma_query_request_timeout_seconds=timeout_seconds,
    )
    return chromadb.HttpClient(host=CHROMA_HOST, port=CHROMA_PORT, settings=settings)


def _load_docx_basic(file_path: str) -> list[Document]:
    try:
        d = docx.Document(file_path)
    except Exception as e:
        raise HTTPException(
            status_code=400, detail=f"Impossible de lire le DOCX: {e}") from e
    texts = []
    for p in d.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text.strip())
    # (option simple; on pourra enrichir avec les tableaux si besoin)
    content = "\n".join(texts).strip()
    if not content:
        return []
    return [Document(page_content=content, metadata={"source": os.path.basename(file_path)})]


def load_docx(file_path: str) -> list[Document]:
    try:
        from unstructured.partition.docx import partition_docx
    except ImportError:  # pragma: no cover - fallback when optional deps missing
        return _load_docx_basic(file_path)

    try:
        elements = partition_docx(filename=file_path, include_metadata=True)
    except Exception:
        logger.warning("partition_docx failed, falling back to basic DOCX loader", exc_info=True)
        return _load_docx_basic(file_path)

    documents: list[Document] = []
    for element in elements:
        text = getattr(element, "text", "") or ""
        text = text.strip()
        if not text:
            continue
        metadata_dict: dict[str, Any] = {}
        metadata_obj = getattr(element, "metadata", None)
        if metadata_obj is not None:
            try:
                raw_meta = metadata_obj.to_dict()
            except AttributeError:
                raw_meta = dict(metadata_obj) if isinstance(metadata_obj, dict) else {}
            for key, value in (raw_meta or {}).items():
                if value in (None, "", [], {}):
                    continue
                metadata_dict[str(key)] = str(value)
        metadata_dict.setdefault("source", os.path.basename(file_path))
        metadata_dict.setdefault(
            "mime_type",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        documents.append(Document(page_content=text, metadata=metadata_dict))

    if not documents:
        return _load_docx_basic(file_path)
    return documents


class TimedOllamaEmbeddings(OllamaEmbeddings):
    """Ollama embeddings client with explicit network timeout."""

    request_timeout: float = OLLAMA_REQUEST_TIMEOUT

    def _process_emb_response(self, input: str) -> list[float]:
        headers = {
            "Content-Type": "application/json",
            **(self.headers or {}),
        }

        try:
            res = requests.post(
                f"{self.base_url}/api/embeddings",
                headers=headers,
                json={"model": self.model, "prompt": input, **self._default_params},
                timeout=self.request_timeout,
            )
        except requests.exceptions.RequestException as e:
            raise ValueError(f"Error raised by inference endpoint: {e}") from e

        if res.status_code != 200:
            raise ValueError(
                f"Error raised by inference API HTTP code: {res.status_code}, {res.text}"
            )
        try:
            t = res.json()
            return t["embedding"]
        except requests.exceptions.JSONDecodeError as e:
            raise ValueError(
                f"Error raised by inference API: {e}.\nResponse: {res.text}"
            ) from e


def load_markdown(file_path: Path) -> list[Document]:
    try:
        raw_text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        raise HTTPException(status_code=400, detail=f"Impossible de lire le fichier Markdown: {exc}") from exc

    text = raw_text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Fichier Markdown vide")

    metadata = {"source": str(file_path), "mime_type": "text/markdown"}
    return [Document(page_content=text, metadata=metadata)]


def _validate_remote_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme.lower() not in URL_SCHEMES_ALLOWED:
        raise HTTPException(
            status_code=400, detail="Schéma d'URL non autorisé")
    if not parsed.hostname:
        raise HTTPException(status_code=400, detail="URL invalide")
    try:
        addr_info = socket.getaddrinfo(parsed.hostname, parsed.port or (
            443 if parsed.scheme == "https" else 80))
    except socket.gaierror as exc:
        raise HTTPException(
            status_code=400, detail=f"Résolution DNS impossible: {exc}") from exc
    for entry in addr_info:
        ip = ipaddress.ip_address(entry[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
            raise HTTPException(
                status_code=400, detail="URL interne non autorisée")


def _download_to_temp(url: str, suffix: str) -> Path:
    _validate_remote_url(url)
    try:
        with requests.get(url, timeout=30, stream=True) as response:
            response.raise_for_status()
            total = 0
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_file:
                for chunk in response.iter_content(chunk_size=8192):
                    if not chunk:
                        continue
                    total += len(chunk)
                    if total > MAX_REMOTE_BYTES:
                        raise HTTPException(
                            status_code=400, detail="Fichier distant trop volumineux")
                    tmp_file.write(chunk)
                return Path(tmp_file.name)
    except HTTPException:
        raise
    except requests.RequestException as exc:
        raise HTTPException(
            status_code=400, detail=f"Téléchargement impossible: {exc}") from exc


def _fetch_remote_text(url: str) -> tuple[str, str]:
    _validate_remote_url(url)
    try:
        with requests.get(url, timeout=30, allow_redirects=True, stream=True) as response:
            if response.history:
                for hop in response.history:
                    _validate_remote_url(hop.url)
            _validate_remote_url(response.url)
            declared_length = response.headers.get("content-length")
            if declared_length and int(declared_length) > MAX_REMOTE_BYTES:
                raise HTTPException(
                    status_code=400, detail="Réponse distante trop volumineuse")
            chunks: list[bytes] = []
            total = 0
            for chunk in response.iter_content(chunk_size=8192):
                if not chunk:
                    continue
                total += len(chunk)
                if total > MAX_REMOTE_BYTES:
                    raise HTTPException(
                        status_code=400, detail="Réponse distante trop volumineuse")
                chunks.append(chunk)
            encoding = response.encoding or "utf-8"
            text = b"".join(chunks).decode(encoding, errors="ignore")
            if not text.strip():
                raise HTTPException(
                    status_code=400, detail="Aucun contenu exploitable sur la page")
            return response.url, text
    except HTTPException:
        raise
    except requests.RequestException as exc:
        raise HTTPException(
            status_code=400, detail=f"Téléchargement impossible: {exc}") from exc


def load_from_url(url: str):
    if url.lower().endswith(".pdf"):
        tmp_path = _download_to_temp(url, suffix=".pdf")
        try:
            return PyPDFLoader(str(tmp_path)).load()
        finally:
            try:
                tmp_path.unlink()
            except OSError:
                pass
    final_url, text = _fetch_remote_text(url)
    soup = BeautifulSoup(text, "html.parser")
    text = soup.get_text("\n", strip=True)
    if not text:
        raise HTTPException(
            status_code=400, detail="Aucun contenu exploitable sur la page")
    return [Document(page_content=text, metadata={"source": final_url})]


def _load_source_documents(req: IngestRequest) -> list[Document]:
    if req.source_type == "url":
        return load_from_url(req.source)
    if req.source_type == "gdrive_folder":
        loader = GoogleDriveLoader(folder_id=req.source, recursive=True)
        return loader.load()
    if req.source_type == "pdf":
        path = _resolve_local_path(req.source, allow_tmp=True)
        return PyPDFLoader(str(path)).load()
    if req.source_type == "docx":
        path = _resolve_local_path(req.source, allow_tmp=True)
        return load_docx(str(path))
    if req.source_type in {"markdown", "md"}:
        path = _resolve_local_path(req.source, allow_tmp=True)
        return load_markdown(path)
    if req.source_type in {"video", "image", "audio"}:
        raise HTTPException(
            status_code=400,
            detail="Ce type de source nécessite le mode multimodal (mode=multimodal).",
        )
    if req.source_type == "auto":
        ext = Path(req.source).suffix.lower()
        if ext == ".pdf":
            path = _resolve_local_path(req.source, allow_tmp=True)
            return PyPDFLoader(str(path)).load()
        if ext == ".docx":
            path = _resolve_local_path(req.source, allow_tmp=True)
            return load_docx(str(path))
        if ext in {".md", ".markdown"}:
            path = _resolve_local_path(req.source, allow_tmp=True)
            return load_markdown(path)
        if ext in {".txt", ".csv", ".json", ".xml", ".html", ".htm"}:
            path = _resolve_local_path(req.source, allow_tmp=True)
            text = path.read_text(encoding="utf-8", errors="ignore").strip()
            if not text:
                return []
            return [Document(page_content=text, metadata={"source": str(path)})]
        # Fallback: try reading as text
        path = _resolve_local_path(req.source, allow_tmp=True)
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            return [Document(page_content=text, metadata={"source": str(path)})]
        return []
    raise HTTPException(status_code=400, detail=f"source_type non géré: {req.source_type}")


def _prepare_chunks_for_chroma(
    req: IngestRequest,
    docs: list[Document],
    splitter: RecursiveCharacterTextSplitter | None = None,
) -> PreparedBatch:
    splitter = splitter or RecursiveCharacterTextSplitter(
        chunk_size=INGEST_CHUNK_SIZE, chunk_overlap=INGEST_CHUNK_OVERLAP
    )
    chunks = splitter.split_documents(docs)
    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict[str, str]] = []
    modality = "text"

    for chunk in chunks:
        text = (chunk.page_content or "").strip()
        if not text:
            continue
        content_hash = get_content_hash(text)
        chunk_modality = (chunk.metadata or {}).get("modality", "text")
        metadata = {
            "sha256": content_hash,
            "source_type": req.source_type,
            "source": req.source,
            "modality": chunk_modality,
        }
        metadata.update(chunk.metadata or {})
        metadata.update(req.metadata_hints or {})
        normalized = normalize_metadata(metadata)

        ids.append(content_hash)
        documents.append(text)
        metadatas.append(normalized)
        modality = normalized.get("modality", modality)

    if not ids:
        modality = "text"
    return PreparedBatch(ids=ids, documents=documents, metadatas=metadatas, modality=modality)


def _prepare_multimodal_chunks(req: IngestRequest, chunks: list[Chunk]) -> PreparedBatch:
    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict[str, str]] = []
    modality_counts: dict[str, int] = {}

    for chunk in chunks:
        text = chunk.as_text() if hasattr(chunk, "as_text") else (chunk.text or "")
        text = (text or "").strip()
        if not text:
            continue
        content_hash = get_content_hash(text)
        chunk_modality = (chunk.modality or "unknown").strip().lower() or "unknown"
        metadata: dict[str, Any] = {
            "sha256": content_hash,
            "source": req.source,
            "source_type": req.source_type,
            "modality": chunk_modality,
        }
        metadata.update(getattr(chunk, "metadata", {}) or {})
        metadata.update(req.metadata_hints or {})
        normalized = normalize_metadata(metadata)

        ids.append(content_hash)
        documents.append(text)
        metadatas.append(normalized)

        key = normalized.get("modality", chunk_modality)
        modality_counts[key] = modality_counts.get(key, 0) + 1

    if modality_counts:
        dominant = max(modality_counts.items(), key=lambda item: item[1])[0]
    else:
        dominant = "unknown"
    return PreparedBatch(ids=ids, documents=documents, metadatas=metadatas, modality=dominant)


_AV_MIME_FALLBACK: dict[str, str] = {
    ".mkv": "video/x-matroska",
    ".webm": "video/webm",
    ".mp4": "video/mp4",
    ".mov": "video/quicktime",
    ".avi": "video/x-msvideo",
    ".flv": "video/x-flv",
    ".mpg": "video/mpeg",
    ".mpeg": "video/mpeg",
    ".ogv": "video/ogg",
    ".mp3": "audio/mpeg",
    ".wav": "audio/wav",
    ".ogg": "audio/ogg",
    ".flac": "audio/flac",
    ".aac": "audio/aac",
    ".m4a": "audio/mp4",
    ".wma": "audio/x-ms-wma",
}


def _guess_mime(filename: str) -> str:
    """Guess MIME type from filename, with fallback for AV formats.

    Uses only the file extension to avoid issues with special Unicode characters
    in filenames that can confuse mimetypes.guess_type.
    """
    ext = Path(filename).suffix.lower()
    fallback = _AV_MIME_FALLBACK.get(ext)
    if fallback:
        return fallback
    mime, _ = mimetypes.guess_type(f"file{ext}")
    return mime or "application/octet-stream"


def _prepare_multimodal_ingest(req: IngestRequest) -> PreparedBatch:
    if not MULTIMODAL_ENABLED:
        raise HTTPException(status_code=400, detail="Multimodal ingest disabled")
    path = _resolve_local_path(req.source)
    mime = _guess_mime(path.name)
    with path.open("rb") as handle:
        chunk_iter = parse_multimodal(
            handle,
            filename=path.name,
            mime=mime,
            timeout_s=MM_PARSER_TIMEOUT,
            max_chars_per_chunk=MM_MAX_CHARS_PER_CHUNK,
            cache_dir=MM_CACHE_DIR,
        )
        chunk_list = list(chunk_iter)
    return _prepare_multimodal_chunks(req, chunk_list)

# --- Endpoint ---


@app.post("/ingest")
def ingest_data(
    req: IngestRequest,
    request: Request,
    mode: str = Query(default="text"),
    section: str = Query(default=""),
    collection: str = Query(default=""),
):
    modality_label = "unknown"
    mode_normalized = (mode or "text").strip().lower() or "text"
    target_col = resolve_collection_name(
        section=section or req.metadata_hints.get("section"),
        collection=collection or req.metadata_hints.get("collection"),
    )

    try:
        _enforce_security(request, req)
    except HTTPException as exc:
        _record_ingest_outcome(req.source_type, modality_label, f"http_{exc.status_code}")
        raise

    try:
        if mode_normalized == "multimodal":
            prepared = _prepare_multimodal_ingest(req)
        elif mode_normalized in {"", "text"}:
            docs = _load_source_documents(req)
            if not docs:
                _record_ingest_metrics(True)
                modality_label = "text"
                _record_ingest_outcome(req.source_type, modality_label, "empty")
                return {"status": "ok", "message": "Aucun document chargé.", "collection": target_col}
            prepared = _prepare_chunks_for_chroma(req, docs)
        else:
            raise HTTPException(status_code=400, detail="Mode d'ingestion non supporté")
        modality_label = prepared.modality or "unknown"
    except HTTPException as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req.source_type, modality_label, f"http_{exc.status_code}")
        raise
    except Exception as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req.source_type, modality_label, "error")
        logger.exception("Erreur de chargement source")
        raise HTTPException(status_code=500, detail="Erreur de chargement du document source") from exc

    if not prepared.ids:
        _record_ingest_metrics(True)
        _record_ingest_outcome(req.source_type, modality_label, "empty")
        return {"status": "ok", "message": "Aucun contenu éligible à l'ingestion.", "collection": target_col}
    prepared, batch_skipped = _dedupe_prepared_batch(prepared)

    try:
        client = get_chroma_client()
        collection_obj = client.get_or_create_collection(
            name=target_col, metadata={"hnsw:space": "cosine"}
        )

        existing = collection_obj.get(ids=prepared.ids) or {}
        existing_ids = set(existing.get("ids", []))

        to_add_idx = [i for i, chunk_id in enumerate(prepared.ids) if chunk_id not in existing_ids]
        if not to_add_idx:
            _record_ingest_metrics(True)
            _record_ingest_outcome(req.source_type, modality_label, "skipped")
            return {
                "status": "ok",
                "added": 0,
                "skipped": len(existing_ids) + batch_skipped,
                "collection": target_col,
            }

        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        docs_to_add = [prepared.documents[i] for i in to_add_idx]
        ids_to_add = [prepared.ids[i] for i in to_add_idx]
        meta_to_add = [prepared.metadatas[i] for i in to_add_idx]
        try:
            embs_to_add = emb.embed_documents(docs_to_add)
        except ValueError as exc:
            message = str(exc)
            if "HTTP code: 404" in message:
                logger.warning(
                    "Ollama embeddings endpoint returned 404 for model '%s'", EMBED_MODEL
                )
                raise HTTPException(
                    status_code=503,
                    detail=f"Modèle d'embedding '{EMBED_MODEL}' non disponible. Vérifiez la configuration.",
                ) from exc
            logger.exception("Embedding provider raised ValueError")
            raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embeddings") from exc
        except Exception as exc:
            logger.exception("Unexpected failure while requesting embeddings")
            raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embeddings") from exc

        meta_mappings = cast(list[Mapping[str, Any]], meta_to_add)
        embeddings_seq = cast(list[Sequence[float]], embs_to_add)
        collection_obj.add(
            documents=docs_to_add,
            ids=ids_to_add,
            metadatas=meta_mappings,
            embeddings=embeddings_seq,
        )
        _record_ingest_metrics(True)
        _record_ingest_outcome(req.source_type, modality_label, "success")
        return {
            "status": "ok",
            "added": len(ids_to_add),
            "skipped": (len(prepared.ids) - len(ids_to_add)) + batch_skipped,
            "collection": target_col,
        }
    except HTTPException as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req.source_type, modality_label, f"http_{exc.status_code}")
        raise
    except Exception as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req.source_type, modality_label, "error")
        logger.exception("Erreur d'ingestion ChromaDB")
        raise HTTPException(status_code=500, detail="Erreur d'indexation dans la base vectorielle") from exc


@app.post("/ingest/upload")
def ingest_upload(
    request: Request,
    file: UploadFile = File(...),  # noqa: B008
    source_type: str = Form(default="auto"),
    mode: str = Form(default="multimodal"),
    metadata: str = Form(default="{}"),
    section: str = Form(default=""),
    collection: str = Form(default=""),
):
    """Ingest a file uploaded via multipart form data.

    This endpoint avoids JSON body encoding issues for filenames
    with special Unicode characters (emojis, accented chars, etc.).
    """
    try:
        _enforce_security(request, None)
    except HTTPException:
        raise

    filename = file.filename or "upload"
    mime = _guess_mime(filename)
    modality_label = "unknown"

    try:
        hints = json.loads(metadata) if metadata else {}
    except (ValueError, TypeError):
        hints = {}

    target_col = resolve_collection_name(
        section=section or hints.get("section"),
        collection=collection or hints.get("collection"),
    )

    # Determine source_type from extension if auto
    if source_type == "auto":
        ext = Path(filename).suffix.lower()
        if ext in {".pdf"}:
            source_type = "pdf"
        elif ext in {".docx"}:
            source_type = "docx"
        elif ext in {".md", ".markdown"}:
            source_type = "markdown"
        elif ext in {".mp4", ".webm", ".mkv", ".mov", ".avi", ".mp3", ".wav", ".ogg", ".flac",
                     ".aac", ".m4a", ".wma"}:
            source_type = "video"
        elif ext in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".webp"}:
            source_type = "image"
        elif ext in {".txt", ".csv", ".json", ".xml", ".html", ".htm"}:
            source_type = "auto"
        else:
            source_type = "auto"

    _validate_upload_mode(source_type, mode)

    # Write uploaded file to temp, then process
    suffix = Path(filename).suffix or ".bin"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False, dir="/tmp") as tmp:
        content = file.file.read()
        file_hash = hashlib.sha256(content).hexdigest()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        ingest_hints = {
            **hints,
            "file_hash": file_hash,
            "original_filename": filename,
        }
        client = get_chroma_client()
        collection_obj = client.get_or_create_collection(
            name=target_col, metadata={"hnsw:space": "cosine"}
        )
        existing_file = collection_obj.get(where={"file_hash": file_hash}) or {}
        existing_file_ids = set(existing_file.get("ids", []))
        if existing_file_ids:
            _record_ingest_metrics(True)
            _record_ingest_outcome(source_type, modality_label, "skipped")
            return {
                "status": "ok",
                "added": 0,
                "skipped": len(existing_file_ids),
                "filename": filename,
                "collection": target_col,
            }
        req = IngestRequest.model_validate({
            "source_type": source_type,
            "source": tmp_path,
            "hints": ingest_hints,
        })
        mode_normalized = (mode or "multimodal").strip().lower()

        if mode_normalized == "multimodal" or source_type in MEDIA_SOURCE_TYPES:
            with open(tmp_path, "rb") as handle:
                chunk_iter = parse_multimodal(
                    handle,
                    filename=filename,
                    mime=mime,
                    timeout_s=MM_PARSER_TIMEOUT,
                    max_chars_per_chunk=MM_MAX_CHARS_PER_CHUNK,
                    cache_dir=MM_CACHE_DIR,
                )
                chunk_list = list(chunk_iter)
            prepared = _prepare_multimodal_chunks(req, chunk_list)
        else:
            docs = _load_source_documents(req)
            if not docs:
                return {"status": "ok", "message": "Aucun document chargé.", "filename": filename, "collection": target_col}
            prepared = _prepare_chunks_for_chroma(req, docs)

        modality_label = prepared.modality or "unknown"

        if not prepared.ids:
            _record_ingest_metrics(True)
            _record_ingest_outcome(source_type, modality_label, "empty")
            return {"status": "ok", "message": "Aucun contenu éligible.", "filename": filename, "collection": target_col}
        prepared, batch_skipped = _dedupe_prepared_batch(prepared)

        existing = collection_obj.get(ids=prepared.ids) or {}
        existing_ids = set(existing.get("ids", []))

        to_add_idx = [i for i, cid in enumerate(prepared.ids) if cid not in existing_ids]

        if not to_add_idx:
            _record_ingest_metrics(True)
            _record_ingest_outcome(source_type, modality_label, "skipped")
            return {
                "status": "ok",
                "added": 0,
                "skipped": len(existing_ids) + batch_skipped,
                "filename": filename,
                "collection": target_col,
            }

        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        docs_to_add = [prepared.documents[i] for i in to_add_idx]
        ids_to_add = [prepared.ids[i] for i in to_add_idx]
        meta_to_add = [prepared.metadatas[i] for i in to_add_idx]
        try:
            embs_to_add = emb.embed_documents(docs_to_add)
        except ValueError as exc:
            message = str(exc)
            if "HTTP code: 404" in message:
                logger.warning(
                    "Ollama embeddings endpoint returned 404 for model '%s'", EMBED_MODEL
                )
                raise HTTPException(
                    status_code=503,
                    detail=f"Modèle d'embedding '{EMBED_MODEL}' non disponible. Vérifiez la configuration.",
                ) from exc
            logger.exception("Embedding provider raised ValueError during upload")
            raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embeddings") from exc
        except Exception as exc:
            logger.exception("Unexpected failure while requesting embeddings during upload")
            raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embeddings") from exc

        meta_mappings = cast(list[Mapping[str, Any]], meta_to_add)
        embeddings_seq = cast(list[Sequence[float]], embs_to_add)
        collection_obj.add(
            documents=docs_to_add,
            ids=ids_to_add,
            metadatas=meta_mappings,
            embeddings=embeddings_seq,
        )
        _record_ingest_metrics(True)
        _record_ingest_outcome(source_type, modality_label, "success")
        return {
            "status": "ok",
            "added": len(ids_to_add),
            "skipped": (len(prepared.ids) - len(ids_to_add)) + batch_skipped,
            "filename": filename,
            "collection": target_col,
        }
    except HTTPException:
        _record_ingest_metrics(False)
        raise
    except Exception as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(source_type, modality_label, "error")
        logger.exception("Erreur d'ingestion upload pour %s", filename)
        raise HTTPException(status_code=500, detail="Erreur d'ingestion du fichier uploadé") from exc
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


@app.get("/health")
def health_check():
    """Public health check — no auth required."""
    return {"status": "healthy"}


@app.get("/metrics")
def metrics(request: Request) -> Response:
    """Prometheus metrics — protected by auth."""
    _enforce_security(request, None)
    if not ingest_metrics.METRICS_ENABLED:
        raise HTTPException(status_code=404, detail="Metrics disabled")
    body = ingest_metrics.generate_latest(METRIC_REGISTRY)
    return Response(body, media_type=CONTENT_TYPE_LATEST)


@app.get("/collections")
def list_collections(request: Request) -> dict[str, Any]:
    """List all ChromaDB collections with document counts."""
    _enforce_security(request, None)
    try:
        client = get_chroma_client()
        collections_raw = client.list_collections()
        result = []
        for col in collections_raw:
            name = col.name if hasattr(col, "name") else str(col)
            try:
                c = client.get_collection(name)
                count = c.count()
            except Exception:
                count = 0
            result.append({"name": name, "count": count})
        return {"collections": result, "total": len(result)}
    except Exception as exc:
        logger.exception("Error listing collections")
        raise HTTPException(status_code=500, detail="Erreur listing collections") from exc


@app.get("/stats/{collection_name}")
def collection_stats(collection_name: str, request: Request) -> dict[str, Any]:
    """Get statistics for a specific collection."""
    _enforce_security(request, None)
    try:
        client = get_chroma_client()
        collection = client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"}
        )
        count = collection.count()
        sample_size = min(count, 100)
        sample = collection.peek(limit=sample_size) if count > 0 else {}
        metadatas = sample.get("metadatas", []) if sample else []

        unique_matieres: set[str] = set()
        unique_niveaux: set[str] = set()
        unique_groupes: set[str] = set()
        unique_types: set[str] = set()
        for m in metadatas:
            if m.get("matiere"):
                unique_matieres.add(m["matiere"])
            if m.get("niveau"):
                unique_niveaux.add(m["niveau"])
            if m.get("groupe"):
                unique_groupes.add(m["groupe"])
            if m.get("type_ressource"):
                unique_types.add(m["type_ressource"])

        return {
            "collection": collection_name,
            "doc_count": count,
            "embed_model": EMBED_MODEL,
            "matieres": sorted(unique_matieres),
            "niveaux": sorted(unique_niveaux),
            "groupes": sorted(unique_groupes),
            "types_ressource": sorted(unique_types),
        }
    except Exception as exc:
        logger.exception("Error getting stats for %s", collection_name)
        raise HTTPException(status_code=500, detail="Erreur stats collection") from exc


class SearchRequest(BaseModel):
    """Search request model with metadata filters."""
    model_config = ConfigDict(extra="ignore")
    q: str
    k: int = Field(default=6, ge=1, le=50)
    section: str | None = None
    collection: str | None = None
    filters: dict[str, Any] | None = None
    include_documents: bool = True
    score_threshold: float | None = Field(default=None, ge=0.0)


@app.post("/search")
def search_kb(payload: SearchRequest, request: Request) -> dict[str, Any]:
    """Semantic search with optional metadata filters."""
    _enforce_security(request, payload)

    target_col = resolve_collection_name(section=payload.section, collection=payload.collection)

    try:
        client = get_chroma_client()
        collection = client.get_or_create_collection(name=target_col, metadata={"hnsw:space": "cosine"})
    except Exception as exc:
        logger.exception("Chroma client error during search")
        raise HTTPException(status_code=500, detail="Erreur de connexion à la base vectorielle") from exc

    try:
        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        q_vec = emb.embed_query(payload.q)
    except ValueError as exc:
        message = str(exc)
        if "HTTP code: 404" in message:
            raise HTTPException(
                status_code=503,
                detail=f"Embedding model '{EMBED_MODEL}' is not available. Pull the model or adjust EMBED_MODEL.",
            ) from exc
        logger.exception("Embedding error during search")
        raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embedding") from exc
    except Exception as exc:
        logger.exception("Unexpected embedding failure during search")
        raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embedding") from exc

    # Build metadata filters for ChromaDB
    where: dict[str, Any] = {}
    if payload.filters:
        conditions = []
        for fk, fv in payload.filters.items():
            if fv is not None and fv != "" and fv != "Tous":
                conditions.append({str(fk): fv})
        if len(conditions) == 1:
            where = conditions[0]
        elif len(conditions) > 1:
            where = {"$and": conditions}

    try:
        n_results = max(1, min(int(payload.k), 50))
        query_kwargs: dict[str, Any] = {"query_embeddings": [q_vec], "n_results": n_results}
        if where:
            query_kwargs["where"] = where
        results = collection.query(**query_kwargs)
    except Exception as exc:
        logger.exception("Chroma query error")
        raise HTTPException(status_code=500, detail="Erreur de recherche") from exc

    documents = results.get("documents", [[]])[0] if results.get("documents") else []
    metadatas = results.get("metadatas", [[]])[0] if results.get("metadatas") else []
    ids = results.get("ids", [[]])[0] if results.get("ids") else []
    distances = results.get("distances", [[]])[0] if results.get("distances") else []

    hits: list[dict[str, Any]] = []
    for idx, doc_id in enumerate(ids):
        distance = distances[idx] if distances and idx < len(distances) else None
        if (
            payload.score_threshold is not None
            and distance is not None
            and float(distance) > payload.score_threshold
        ):
            continue
        item: dict[str, Any] = {"id": doc_id, "metadata": metadatas[idx] if idx < len(metadatas) else {}}
        if payload.include_documents and idx < len(documents):
            item["document"] = documents[idx]
        if distance is not None:
            item["score"] = distance
        hits.append(item)

    _record_ingest_outcome("search", "text", "success")
    return {
        "query": payload.q,
        "collection": target_col,
        "k": n_results,
        "returned": len(hits),
        "filters_applied": where,
        "score_threshold": payload.score_threshold,
        "hits": hits,
    }


class RagQueryFilters(BaseModel):
    domain: str | None = None
    document_id: str | None = None
    tags: list[str] | None = None
    metadata: dict[str, Any] | None = None


class RagQuery(BaseModel):
    query: str
    filters: RagQueryFilters | None = None
    top_k: int = Field(default=6, ge=1, le=50)
    collection: str = Field(default=COLLECTION_NAME)


@app.post("/rag/query")
def rag_query(payload: RagQuery, request: Request) -> dict[str, Any]:
    """RAG query with domain/tags/metadata filters."""
    _enforce_security(request, payload)

    try:
        client = get_chroma_client()
        collection = client.get_or_create_collection(name=payload.collection, metadata={"hnsw:space": "cosine"})
    except Exception as exc:
        logger.exception("Chroma client error during rag/query")
        raise HTTPException(status_code=500, detail="Erreur de connexion à la base vectorielle") from exc

    try:
        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        q_vec = emb.embed_query(payload.query)
    except ValueError as exc:
        message = str(exc)
        if "HTTP code: 404" in message:
            raise HTTPException(
                status_code=503,
                detail=f"Embedding model '{EMBED_MODEL}' is not available. Pull the model or adjust EMBED_MODEL.",
            ) from exc
        logger.exception("Embedding error during rag/query")
        raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embedding") from exc
    except Exception as exc:
        logger.exception("Unexpected embedding failure during rag/query")
        raise HTTPException(status_code=500, detail="Erreur lors du calcul d'embedding") from exc

    where: dict[str, Any] = {}
    if payload.filters:
        f = payload.filters
        if f.domain:
            where["domain"] = f.domain
        if f.document_id:
            where["document_id"] = f.document_id
        if f.tags:
            where["tags"] = {"$in": f.tags}
        if f.metadata:
            for k, v in f.metadata.items():
                if v is not None and v != "":
                    where[str(k)] = v

    try:
        n_results = max(1, min(int(payload.top_k), 50))
        query_kwargs: dict[str, Any] = {"query_embeddings": [q_vec], "n_results": n_results}
        if where:
            query_kwargs["where"] = where
        results = collection.query(**query_kwargs)
    except Exception as exc:
        logger.exception("Chroma query error during rag/query")
        raise HTTPException(status_code=500, detail="Erreur de recherche") from exc

    documents = results.get("documents", [[]])[0] if results.get("documents") else []
    metadatas = results.get("metadatas", [[]])[0] if results.get("metadatas") else []
    ids = results.get("ids", [[]])[0] if results.get("ids") else []
    distances = results.get("distances", [[]])[0] if results.get("distances") else []

    hits: list[dict[str, Any]] = []
    for idx, doc_id in enumerate(ids):
        item: dict[str, Any] = {"id": doc_id, "metadata": metadatas[idx] if idx < len(metadatas) else {}}
        if idx < len(documents):
            item["document"] = documents[idx]
        if distances and idx < len(distances) and distances[idx] is not None:
            item["score"] = distances[idx]
        hits.append(item)

    return {
        "query": payload.query,
        "collection": payload.collection,
        "k": n_results,
        "filters": where,
        "hits": hits,
    }
